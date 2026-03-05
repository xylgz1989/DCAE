"""
DCAE (Disciplined Consensus-Driven Agentic Engineering) Framework
Requirements Export Module

This module implements the export functionality for requirements as specified in
Story 2.5: Export Share Requirements.

As a user,
I want to export or share requirements documents,
so that I can collaborate with team members and stakeholders outside the DCAE system.
"""

import os
import json
import yaml
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime
import csv

# Try to import optional PDF generation libraries
try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class RequirementsExporter:
    """Handles exporting requirements in various formats."""

    def __init__(self, requirements: Dict[str, Any]):
        """
        Initialize the exporter with requirements data.

        Args:
            requirements: Dictionary containing requirements data
        """
        self.requirements = requirements

    def export_to_txt(self, output_path: Path) -> bool:
        """
        Export requirements to plain text format.

        Args:
            output_path: Path to save the exported file

        Returns:
            Boolean indicating success
        """
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(f"Requirements Document: {self.requirements.get('project_name', 'Untitled')}\n")
                f.write("=" * 50 + "\n\n")

                # Write project description
                description = self.requirements.get('description', '')
                if description:
                    f.write(f"Project Description:\n{description}\n\n")

                # Write functional requirements
                func_reqs = self.requirements.get('functional_requirements', [])
                if func_reqs:
                    f.write("Functional Requirements:\n")
                    f.write("-" * 30 + "\n")
                    for i, req in enumerate(func_reqs, 1):
                        if isinstance(req, dict):
                            req_id = req.get('id', f'FR{i:03d}')
                            req_desc = req.get('description', 'No description')
                            req_priority = req.get('priority', 'N/A')
                            f.write(f"{req_id}: {req_desc}\n")
                            f.write(f"  Priority: {req_priority}\n\n")
                        else:
                            f.write(f"FR{i:03d}: {req}\n\n")

                # Write non-functional requirements
                non_func_reqs = self.requirements.get('non_functional_requirements', [])
                if non_func_reqs:
                    f.write("Non-Functional Requirements:\n")
                    f.write("-" * 35 + "\n")
                    for i, req in enumerate(non_func_reqs, 1):
                        if isinstance(req, dict):
                            req_id = req.get('id', f'NFR{i:03d}')
                            req_desc = req.get('description', 'No description')
                            req_category = req.get('category', 'N/A')
                            req_priority = req.get('priority', 'N/A')
                            f.write(f"{req_id}: {req_desc}\n")
                            f.write(f"  Category: {req_category}\n")
                            f.write(f"  Priority: {req_priority}\n\n")
                        else:
                            f.write(f"NFR{i:03d}: {req}\n\n")

                # Write constraints
                constraints = self.requirements.get('constraints', [])
                if constraints:
                    f.write("Constraints:\n")
                    f.write("-" * 15 + "\n")
                    for i, req in enumerate(constraints, 1):
                        if isinstance(req, dict):
                            req_id = req.get('id', f'C{i:03d}')
                            req_desc = req.get('description', 'No description')
                            f.write(f"{req_id}: {req_desc}\n\n")
                        else:
                            f.write(f"C{i:03d}: {req}\n\n")

                # Write assumptions
                assumptions = self.requirements.get('assumptions', [])
                if assumptions:
                    f.write("Assumptions:\n")
                    f.write("-" * 15 + "\n")
                    for i, req in enumerate(assumptions, 1):
                        if isinstance(req, dict):
                            req_id = req.get('id', f'A{i:03d}')
                            req_desc = req.get('description', 'No description')
                            f.write(f"{req_id}: {req_desc}\n\n")
                        else:
                            f.write(f"A{i:03d}: {req}\n\n")

                # Write acceptance criteria
                acceptance = self.requirements.get('acceptance_criteria', [])
                if acceptance:
                    f.write("Acceptance Criteria:\n")
                    f.write("-" * 22 + "\n")
                    for i, req in enumerate(acceptance, 1):
                        if isinstance(req, dict):
                            req_id = req.get('id', f'AC{i:03d}')
                            req_desc = req.get('description', 'No description')
                            f.write(f"{req_id}: {req_desc}\n\n")
                        else:
                            f.write(f"AC{i:03d}: {req}\n\n")

                # Write metadata
                metadata = self.requirements.get('metadata', {})
                if metadata:
                    f.write("Metadata:\n")
                    f.write("-" * 10 + "\n")
                    for key, value in metadata.items():
                        f.write(f"{key}: {value}\n")

            return True
        except Exception as e:
            print(f"Error exporting to TXT: {e}")
            return False

    def export_to_csv(self, output_path: Path) -> bool:
        """
        Export requirements to CSV format.

        Args:
            output_path: Path to save the exported file

        Returns:
            Boolean indicating success
        """
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)

            with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
                fieldnames = ['Type', 'ID', 'Description', 'Category', 'Priority', 'Extra Info']
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)

                writer.writeheader()

                # Write functional requirements
                func_reqs = self.requirements.get('functional_requirements', [])
                for i, req in enumerate(func_reqs, 1):
                    if isinstance(req, dict):
                        writer.writerow({
                            'Type': 'Functional',
                            'ID': req.get('id', f'FR{i:03d}'),
                            'Description': req.get('description', ''),
                            'Category': '',
                            'Priority': req.get('priority', 'N/A'),
                            'Extra Info': ''
                        })
                    else:
                        writer.writerow({
                            'Type': 'Functional',
                            'ID': f'FR{i:03d}',
                            'Description': req,
                            'Category': '',
                            'Priority': 'N/A',
                            'Extra Info': ''
                        })

                # Write non-functional requirements
                non_func_reqs = self.requirements.get('non_functional_requirements', [])
                for i, req in enumerate(non_func_reqs, 1):
                    if isinstance(req, dict):
                        writer.writerow({
                            'Type': 'Non-Functional',
                            'ID': req.get('id', f'NFR{i:03d}'),
                            'Description': req.get('description', ''),
                            'Category': req.get('category', 'N/A'),
                            'Priority': req.get('priority', 'N/A'),
                            'Extra Info': ''
                        })
                    else:
                        writer.writerow({
                            'Type': 'Non-Functional',
                            'ID': f'NFR{i:03d}',
                            'Description': req,
                            'Category': 'N/A',
                            'Priority': 'N/A',
                            'Extra Info': ''
                        })

                # Write constraints
                constraints = self.requirements.get('constraints', [])
                for i, req in enumerate(constraints, 1):
                    if isinstance(req, dict):
                        writer.writerow({
                            'Type': 'Constraint',
                            'ID': req.get('id', f'C{i:03d}'),
                            'Description': req.get('description', ''),
                            'Category': '',
                            'Priority': 'N/A',
                            'Extra Info': ''
                        })
                    else:
                        writer.writerow({
                            'Type': 'Constraint',
                            'ID': f'C{i:03d}',
                            'Description': req,
                            'Category': '',
                            'Priority': 'N/A',
                            'Extra Info': ''
                        })

                # Write assumptions
                assumptions = self.requirements.get('assumptions', [])
                for i, req in enumerate(assumptions, 1):
                    if isinstance(req, dict):
                        writer.writerow({
                            'Type': 'Assumption',
                            'ID': req.get('id', f'A{i:03d}'),
                            'Description': req.get('description', ''),
                            'Category': '',
                            'Priority': 'N/A',
                            'Extra Info': ''
                        })
                    else:
                        writer.writerow({
                            'Type': 'Assumption',
                            'ID': f'A{i:03d}',
                            'Description': req,
                            'Category': '',
                            'Priority': 'N/A',
                            'Extra Info': ''
                        })

                # Write acceptance criteria
                acceptance = self.requirements.get('acceptance_criteria', [])
                for i, req in enumerate(acceptance, 1):
                    if isinstance(req, dict):
                        writer.writerow({
                            'Type': 'Acceptance',
                            'ID': req.get('id', f'AC{i:03d}'),
                            'Description': req.get('description', ''),
                            'Category': '',
                            'Priority': 'N/A',
                            'Extra Info': ''
                        })
                    else:
                        writer.writerow({
                            'Type': 'Acceptance',
                            'ID': f'AC{i:03d}',
                            'Description': req,
                            'Category': '',
                            'Priority': 'N/A',
                            'Extra Info': ''
                        })

            return True
        except Exception as e:
            print(f"Error exporting to CSV: {e}")
            return False

    def export_to_pdf(self, output_path: Path) -> bool:
        """
        Export requirements to PDF format.

        Args:
            output_path: Path to save the exported file

        Returns:
            Boolean indicating success
        """
        if not PDF_AVAILABLE:
            print("FPDF library not available. Install with 'pip install fpdf2'")
            return False

        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)

            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()

            # Title
            pdf.set_font("Arial", 'B', 16)
            pdf.cell(0, 10, f"Requirements Document: {self.requirements.get('project_name', 'Untitled')}", ln=True, align='C')
            pdf.ln(10)

            # Project description
            description = self.requirements.get('description', '')
            if description:
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 10, "Project Description:", ln=True)
                pdf.set_font("Arial", '', 10)
                # Adding description with proper line breaks
                for line in description.split('\n'):
                    pdf.cell(0, 8, line, ln=True)
                pdf.ln(5)

            # Functional requirements
            func_reqs = self.requirements.get('functional_requirements', [])
            if func_reqs:
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 10, "Functional Requirements:", ln=True)
                pdf.set_font("Arial", '', 10)
                for i, req in enumerate(func_reqs, 1):
                    if isinstance(req, dict):
                        req_id = req.get('id', f'FR{i:03d}')
                        req_desc = req.get('description', 'No description')
                        req_priority = req.get('priority', 'N/A')
                        pdf.cell(0, 8, f"{req_id}: {req_desc}", ln=True)
                        pdf.cell(0, 8, f"  Priority: {req_priority}", ln=True)
                    else:
                        pdf.cell(0, 8, f"FR{i:03d}: {req}", ln=True)
                    pdf.ln(2)

            # Non-functional requirements
            non_func_reqs = self.requirements.get('non_functional_requirements', [])
            if non_func_reqs:
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 10, "Non-Functional Requirements:", ln=True)
                pdf.set_font("Arial", '', 10)
                for i, req in enumerate(non_func_reqs, 1):
                    if isinstance(req, dict):
                        req_id = req.get('id', f'NFR{i:03d}')
                        req_desc = req.get('description', 'No description')
                        req_category = req.get('category', 'N/A')
                        req_priority = req.get('priority', 'N/A')
                        pdf.cell(0, 8, f"{req_id}: {req_desc}", ln=True)
                        pdf.cell(0, 8, f"  Category: {req_category}, Priority: {req_priority}", ln=True)
                    else:
                        pdf.cell(0, 8, f"NFR{i:03d}: {req}", ln=True)
                    pdf.ln(2)

            # Constraints
            constraints = self.requirements.get('constraints', [])
            if constraints:
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 10, "Constraints:", ln=True)
                pdf.set_font("Arial", '', 10)
                for i, req in enumerate(constraints, 1):
                    if isinstance(req, dict):
                        req_id = req.get('id', f'C{i:03d}')
                        req_desc = req.get('description', 'No description')
                        pdf.cell(0, 8, f"{req_id}: {req_desc}", ln=True)
                    else:
                        pdf.cell(0, 8, f"C{i:03d}: {req}", ln=True)
                    pdf.ln(2)

            # Assumptions
            assumptions = self.requirements.get('assumptions', [])
            if assumptions:
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 10, "Assumptions:", ln=True)
                pdf.set_font("Arial", '', 10)
                for i, req in enumerate(assumptions, 1):
                    if isinstance(req, dict):
                        req_id = req.get('id', f'A{i:03d}')
                        req_desc = req.get('description', 'No description')
                        pdf.cell(0, 8, f"{req_id}: {req_desc}", ln=True)
                    else:
                        pdf.cell(0, 8, f"A{i:03d}: {req}", ln=True)
                    pdf.ln(2)

            # Acceptance criteria
            acceptance = self.requirements.get('acceptance_criteria', [])
            if acceptance:
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 10, "Acceptance Criteria:", ln=True)
                pdf.set_font("Arial", '', 10)
                for i, req in enumerate(acceptance, 1):
                    if isinstance(req, dict):
                        req_id = req.get('id', f'AC{i:03d}')
                        req_desc = req.get('description', 'No description')
                        pdf.cell(0, 8, f"{req_id}: {req_desc}", ln=True)
                    else:
                        pdf.cell(0, 8, f"AC{i:03d}: {req}", ln=True)
                    pdf.ln(2)

            # Metadata
            metadata = self.requirements.get('metadata', {})
            if metadata:
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 10, "Metadata:", ln=True)
                pdf.set_font("Arial", '', 10)
                for key, value in metadata.items():
                    pdf.cell(0, 8, f"{key}: {value}", ln=True)

            pdf.output(str(output_path))
            return True
        except Exception as e:
            print(f"Error exporting to PDF: {e}")
            return False

    def export_to_docx(self, output_path: Path) -> bool:
        """
        Export requirements to DOCX format.

        Args:
            output_path: Path to save the exported file

        Returns:
            Boolean indicating success
        """
        if not DOCX_AVAILABLE:
            print("python-docx library not available. Install with 'pip install python-docx'")
            return False

        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)

            doc = Document()

            # Add title
            doc.add_heading(f'Requirements Document: {self.requirements.get("project_name", "Untitled")}', 0)

            # Add project description
            description = self.requirements.get('description', '')
            if description:
                doc.add_heading('Project Description', level=1)
                doc.add_paragraph(description)

            # Add functional requirements
            func_reqs = self.requirements.get('functional_requirements', [])
            if func_reqs:
                doc.add_heading('Functional Requirements', level=1)
                for i, req in enumerate(func_reqs, 1):
                    if isinstance(req, dict):
                        req_id = req.get('id', f'FR{i:03d}')
                        req_desc = req.get('description', 'No description')
                        req_priority = req.get('priority', 'N/A')
                        para = doc.add_paragraph(style='List Number')
                        para.add_run(f'{req_id}: {req_desc}').bold = True
                        doc.add_paragraph(f'Priority: {req_priority}')
                    else:
                        para = doc.add_paragraph(style='List Number')
                        para.add_run(f'FR{i:03d}: {req}').bold = True

            # Add non-functional requirements
            non_func_reqs = self.requirements.get('non_functional_requirements', [])
            if non_func_reqs:
                doc.add_heading('Non-Functional Requirements', level=1)
                for i, req in enumerate(non_func_reqs, 1):
                    if isinstance(req, dict):
                        req_id = req.get('id', f'NFR{i:03d}')
                        req_desc = req.get('description', 'No description')
                        req_category = req.get('category', 'N/A')
                        req_priority = req.get('priority', 'N/A')
                        para = doc.add_paragraph(style='List Number')
                        para.add_run(f'{req_id}: {req_desc}').bold = True
                        doc.add_paragraph(f'Category: {req_category}, Priority: {req_priority}')
                    else:
                        para = doc.add_paragraph(style='List Number')
                        para.add_run(f'NFR{i:03d}: {req}').bold = True

            # Add constraints
            constraints = self.requirements.get('constraints', [])
            if constraints:
                doc.add_heading('Constraints', level=1)
                for i, req in enumerate(constraints, 1):
                    if isinstance(req, dict):
                        req_id = req.get('id', f'C{i:03d}')
                        req_desc = req.get('description', 'No description')
                        para = doc.add_paragraph(style='List Number')
                        para.add_run(f'{req_id}: {req_desc}').bold = True
                    else:
                        para = doc.add_paragraph(style='List Number')
                        para.add_run(f'C{i:03d}: {req}').bold = True

            # Add assumptions
            assumptions = self.requirements.get('assumptions', [])
            if assumptions:
                doc.add_heading('Assumptions', level=1)
                for i, req in enumerate(assumptions, 1):
                    if isinstance(req, dict):
                        req_id = req.get('id', f'A{i:03d}')
                        req_desc = req.get('description', 'No description')
                        para = doc.add_paragraph(style='List Number')
                        para.add_run(f'{req_id}: {req_desc}').bold = True
                    else:
                        para = doc.add_paragraph(style='List Number')
                        para.add_run(f'A{i:03d}: {req}').bold = True

            # Add acceptance criteria
            acceptance = self.requirements.get('acceptance_criteria', [])
            if acceptance:
                doc.add_heading('Acceptance Criteria', level=1)
                for i, req in enumerate(acceptance, 1):
                    if isinstance(req, dict):
                        req_id = req.get('id', f'AC{i:03d}')
                        req_desc = req.get('description', 'No description')
                        para = doc.add_paragraph(style='List Number')
                        para.add_run(f'{req_id}: {req_desc}').bold = True
                    else:
                        para = doc.add_paragraph(style='List Number')
                        para.add_run(f'AC{i:03d}: {req}').bold = True

            # Add metadata
            metadata = self.requirements.get('metadata', {})
            if metadata:
                doc.add_heading('Metadata', level=1)
                for key, value in metadata.items():
                    doc.add_paragraph(f'{key}: {value}')

            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Error exporting to DOCX: {e}")
            return False

    def export_by_format(self, output_path: Path, format_type: str) -> bool:
        """
        Export requirements to a specific format.

        Args:
            output_path: Path to save the exported file
            format_type: Format to export ('txt', 'csv', 'pdf', 'docx', 'json', 'yaml')

        Returns:
            Boolean indicating success
        """
        format_map = {
            'txt': self.export_to_txt,
            'csv': self.export_to_csv,
            'pdf': self.export_to_pdf,
            'docx': self.export_to_docx,
            'json': lambda path: self._export_to_json(path),
            'yaml': lambda path: self._export_to_yaml(path)
        }

        if format_type.lower() in format_map:
            return format_map[format_type.lower()](output_path)
        else:
            print(f"Unsupported format: {format_type}")
            return False

    def _export_to_json(self, output_path: Path) -> bool:
        """
        Export requirements to JSON format.

        Args:
            output_path: Path to save the exported file

        Returns:
            Boolean indicating success
        """
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)

            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(self.requirements, f, indent=2, ensure_ascii=False, default=str)

            return True
        except Exception as e:
            print(f"Error exporting to JSON: {e}")
            return False

    def _export_to_yaml(self, output_path: Path) -> bool:
        """
        Export requirements to YAML format.

        Args:
            output_path: Path to save the exported file

        Returns:
            Boolean indicating success
        """
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)

            with open(output_path, 'w', encoding='utf-8') as f:
                yaml.dump(self.requirements, f, default_flow_style=False, allow_unicode=True, indent=2)

            return True
        except Exception as e:
            print(f"Error exporting to YAML: {e}")
            return False


def export_requirements(requirements_path: Path, output_path: Path, format_type: str) -> bool:
    """
    Convenience function to export requirements from a file to a specified format.

    Args:
        requirements_path: Path to the input requirements file
        output_path: Path to save the exported file
        format_type: Format to export ('txt', 'csv', 'pdf', 'docx', 'json', 'yaml')

    Returns:
        Boolean indicating success
    """
    from .requirements import load_requirements  # Import locally to avoid circular imports

    requirements = load_requirements(requirements_path)
    if not requirements:
        print(f"No requirements found at {requirements_path}")
        return False

    exporter = RequirementsExporter(requirements)
    return exporter.export_by_format(output_path, format_type)


def create_shareable_link(requirements_path: Path, expiration_hours: int = 24) -> str:
    """
    Create a shareable link for requirements.

    NOTE: This is a placeholder implementation. A real implementation would require
    server-side infrastructure for secure file sharing.

    Args:
        requirements_path: Path to the requirements file
        expiration_hours: Number of hours until the link expires (default 24)

    Returns:
        String containing the shareable link
    """
    import hashlib
    import time

    # This is a placeholder - in a real system, this would upload to a secure server
    # and return a unique URL with access controls and expiration
    file_content = requirements_path.read_text(encoding='utf-8')
    hash_content = hashlib.sha256(file_content.encode()).hexdigest()

    # Generate a unique identifier based on the content and timestamp
    timestamp = str(int(time.time()))
    unique_id = f"{hash_content[:12]}_{timestamp}_{expiration_hours}h"

    # Placeholder URL - would point to an actual secure sharing endpoint
    share_url = f"https://dcae.example.com/share/{unique_id}"

    print(f"Shareable link created: {share_url}")
    print(f"Link will expire in {expiration_hours} hours")

    return share_url